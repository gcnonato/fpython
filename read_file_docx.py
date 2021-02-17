import os
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup as bs

# import textract
# import win32com.client

import olefile



homepath = os.path.expanduser(os.getenv("USERPROFILE"))
desktoppath = "Desktop"
filename = 'temporario.doc'
local_filename = os.path.join(homepath, desktoppath, filename)

print()
assert olefile.isOleFile(local_filename)

with olefile.OleFileIO(local_filename) as ole:
    print(ole)


# word = win32com.client.Dispatch('Word.Application')
# word.visible = False
# wb = word.Documents.Open(local_filename)
# doc = word.ActiveDocument
# print(doc.Range().Text)

# text = textract.process(local_filename)

with open(local_filename, encoding='ISO-8859-1') as _f:
    print(_f.read())
try:
    soup = bs(open(local_filename).read())
    print(soup)
    [s.extract() for s in soup(['style', 'script'])]
    tmpText = soup.get_text()
    text = "".join("".join(tmpText.split('\t')).split('\n')).encode('utf-8').strip()
    print(text)
except UnicodeDecodeError as err:
    print(err)
# document = Document('../temporario.doc')

def create_docx(document):
    document.add_heading('Título do Documento', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    document.add_picture('monty-truth.jpg', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save('../demo.docx')
